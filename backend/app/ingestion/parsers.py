"""Extracción de texto desde distintos formatos de documento.

Cada parser devuelve una lista de `ExtractedSegment`, preservando la unidad
lógica del documento (página de PDF, hoja de Excel, etc.) para poder citar
la fuente exacta más adelante.
"""
import csv
import io
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from ebooklib import ITEM_DOCUMENT
from ebooklib.epub import EpubNav, read_epub
from pypdf import PdfReader


class UnsupportedFileTypeError(Exception):
    """Se lanza cuando la extensión del archivo no está soportada."""


class DocumentParseError(Exception):
    """Se lanza cuando un archivo no pudo ser parseado (corrupto, vacío, etc.)."""


@dataclass
class ExtractedSegment:
    """Un segmento de texto extraído de un documento, con su ubicación de origen."""

    text: str
    location: str  # ej. "página 3", "hoja Ventas", "línea 1-50"


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".csv", ".epub"}


def parse_document(file_path: Path) -> list[ExtractedSegment]:
    """Despacha al parser adecuado según la extensión del archivo."""
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _parse_pdf(file_path)
        if suffix == ".docx":
            return _parse_docx(file_path)
        if suffix in (".xlsx", ".xls"):
            return _parse_xlsx(file_path)
        if suffix == ".txt":
            return _parse_txt(file_path)
        if suffix == ".csv":
            return _parse_csv(file_path)
        if suffix == ".epub":
            return _parse_epub(file_path)
    except UnsupportedFileTypeError:
        raise
    except Exception as exc:  # noqa: BLE001 - queremos capturar cualquier fallo de parsing
        raise DocumentParseError(f"No se pudo procesar '{file_path.name}': {exc}") from exc

    raise UnsupportedFileTypeError(
        f"Extensión '{suffix}' no soportada. Formatos válidos: {sorted(SUPPORTED_EXTENSIONS)}"
    )


def _parse_pdf(file_path: Path) -> list[ExtractedSegment]:
    reader = PdfReader(str(file_path))
    segments = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            segments.append(ExtractedSegment(text=text, location=f"página {page_num}"))
    if not segments:
        raise DocumentParseError("El PDF no contiene texto extraíble (¿está escaneado como imagen?).")
    return segments


def _parse_docx(file_path: Path) -> list[ExtractedSegment]:
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Incluir también el contenido de tablas, que python-docx no expone en .paragraphs
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)

    full_text = "\n".join(paragraphs + table_texts)
    if not full_text.strip():
        raise DocumentParseError("El documento DOCX está vacío.")
    return [ExtractedSegment(text=full_text, location="documento completo")]


def _parse_xlsx(file_path: Path) -> list[ExtractedSegment]:
    sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
    segments = []
    for sheet_name, df in sheets.items():
        df = df.fillna("")
        if df.empty:
            continue
        lines = [" | ".join(str(col) for col in df.columns)]
        for _, row in df.iterrows():
            lines.append(" | ".join(str(v) for v in row.values))
        segments.append(ExtractedSegment(text="\n".join(lines), location=f"hoja '{sheet_name}'"))
    if not segments:
        raise DocumentParseError("El archivo Excel no contiene datos.")
    return segments


def _parse_txt(file_path: Path) -> list[ExtractedSegment]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        raise DocumentParseError("El archivo TXT está vacío.")
    return [ExtractedSegment(text=text, location="documento completo")]


def _parse_csv(file_path: Path) -> list[ExtractedSegment]:
    with open(file_path, encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        raise DocumentParseError("El archivo CSV está vacío.")
    text = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
    return [ExtractedSegment(text=text, location="documento completo")]


def _parse_epub(file_path: Path) -> list[ExtractedSegment]:
    book = read_epub(str(file_path), options={"ignore_ncx": True})

    segments = []
    index = 0
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        # El documento de navegación (tabla de contenidos) no aporta contenido real
        if isinstance(item, EpubNav):
            continue

        index += 1
        soup = BeautifulSoup(item.get_content(), "html.parser")

        # Título del capítulo si existe, si no se usa un índice secuencial
        heading = soup.find(["h1", "h2", "h3"])
        chapter_title = heading.get_text(strip=True) if heading else None
        location = f"capítulo {index}" + (f" — {chapter_title}" if chapter_title else "")

        text = soup.get_text(separator="\n").strip()
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        if text:
            segments.append(ExtractedSegment(text=text, location=location))

    if not segments:
        raise DocumentParseError("El EPUB no contiene texto extraíble.")
    return segments
